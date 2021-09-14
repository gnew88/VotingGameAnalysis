import pandas as pd
from datetime import date
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


class reportGenerator:
    def __init__(self):
        self.document = Document()
        self.document.styles['Normal'].font.name = u'微軟正黑體'
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微軟正黑體')
    
    ######### put table in our document ##########
    def add_table(self, table):
        t = self.document.add_table(table.shape[0]+1, table.shape[1])
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style='Table Grid' 

        # add the header rows.
        for j in range(table.shape[-1]):
            t.cell(0,j).text = table.columns[j]

        # add the rest of the data frame
        for i in range(table.shape[0]):
            for j in range(table.shape[-1]):
                t.cell(i+1,j).text = str(table.values[i,j])

        for row in t.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(10)

    def basic_info(self, files, control_amount, treatment_amount):
        ######## title ########
        self.document.add_heading('【Voting Game】資料分析報告', level = 0)

        ######## basic information ########
        self.document.add_heading('基本資訊', level = 1)

        # today's date 
        date_p = self.document.add_paragraph()
        today = date.today()
        day = today.strftime("%B %d, %Y")
        date_p.add_run('日期：').bold = True
        date_p.add_run(day)

        # the files we have used
        file_p = self.document.add_paragraph()
        file_p.add_run('資料集：').bold = True
        for idx in range(len(files)):
            if idx != len(files) - 1:
                file_p.add_run(files[idx] + ", ")
            else:
                file_p.add_run(files[idx])

        # Number of observations
        ct_p = self.document.add_paragraph()
        ct_p.add_run('觀測值數量：').bold = True
        ct_p.add_run(f"{control_amount + treatment_amount} （全部）, {control_amount} （控制組）, {treatment_amount} （實驗組）")
    
    def participants_distribution(self, distribution_table):
        ######### Distribution of Observations ########
        self.document.add_heading('資料分布', level = 1)
        distriInfo_p = self.document.add_paragraph()
        distriInfo_p.add_run('- 檢視控制組與實驗組的資料筆')
        distriInfo_p.add_run('\n- 檢視在兩場遊戲當中之大小隊分組情形')
        self.add_table(distribution_table)
    
    def data_analysis(self, paired_table, largePq_int_all, largePq_int_sin, largePq_no_all, largePq_no_sin,
                      ct_wtp, sum_reg, sum_f, max_reg, max_f, min_reg, min_f):
        ######### Data Analysis ########
        
        ##### paired t test
        self.document.add_heading('【C】有門檻 wtp v.s. 沒門檻 wtp', level = 1)
        pairedT_p = self.document.add_paragraph()
        pairedT_p.add_run('- 僅使用控制組的資料集')
        pairedT_p.add_run('\n- 透過成對 t 檢定檢驗在不同隊伍分配情形下，門檻的有無是否造成 wtp 的差異')
        self.add_table(paired_table)
        
        ##### The interaction - is_large, is_pq
        space = self.document.add_paragraph()
        space.add_run(' ')        

        self.document.add_heading('【C】檢視隊伍大小與有無門檻的選舉之間是否有交互作用', level = 1)
        largePq_p = self.document.add_paragraph()
        largePq_p.add_run('- 僅使用控制組的資料集')
        largePq_p.add_run('\n- 模型設定: reg wtp is_large is_pq is_large*is_pq')
        largePq_p.add_run('\n- 其中 wtp 代表該場遊戲的願付價格，is_large 則是其是否被分配於大組，is_pq 則是該場遊戲有無門檻設置')
        largePq_p.add_run('\n- 我們同時也建立沒有加入交乘項的模型')
        
        space = self.document.add_paragraph()
        space.add_run(' ')

        self.document.add_heading('有加入交乘項 - 全體玩家', level = 2)
        self.add_table(largePq_int_all)
        
        space = self.document.add_paragraph()
        space.add_run(' ')

        self.document.add_heading('有加入交乘項 - 單一玩家', level = 2)
        largePq_sin_p = self.document.add_paragraph()
        largePq_sin_p.add_run('- 針對個別玩家建立迴歸模型')
        largePq_sin_p.add_run('\n- 分別統計滿足變數係數正負向組合的人數')
        self.add_table(largePq_int_sin)
        
        space = self.document.add_paragraph()
        space.add_run(' ')

        self.document.add_heading('沒有加入交乘項 - 所有玩家', level = 2)
        self.add_table(largePq_no_all)
        
        space = self.document.add_paragraph()
        space.add_run(' ')

        self.document.add_heading('沒有加入交乘項 - 單一玩家', level = 2)
        self.add_table(largePq_no_sin)
        
        ######### Control v.s. Treatment
        space = self.document.add_paragraph()
        space.add_run(' ')
        self.document.add_heading('控制組 wtp v.s. 實驗組 wtp', level = 1)
        ct_p = self.document.add_paragraph()
        ct_p.add_run('- 使用控制組與實驗組的資料')
        ct_p.add_run('\n- 控制組為兩個遊戲分開進行，實驗組則是合一')
        ct_p.add_run('\n- 控制組的 wtp 計算方式為兩個遊戲 wtp 之總和，而實驗組的則是原始 wtp 數據')
        ct_p.add_run('\n- 我們使用未成對 t 檢定檢驗控制組與實驗組之 wtp 是否有顯著差異')
        self.add_table(ct_wtp)
        
        ######### interaction: treatment, sb_condition
        space = self.document.add_paragraph()
        space.add_run(' ')
        
        self.document.add_heading('檢視隊伍分配情形與實驗組效果之間是否有交互作用', level = 1)
        tr_sb_p = self.document.add_paragraph()
        tr_sb_p.add_run('- 使用控制組與實驗組的資料')
        tr_sb_p.add_run('\n- 模型設定：reg wtp treatment sb bs bb sb*treatment bs_treatment bb_treatment')
        tr_sb_p.add_run('\n- 其中 treatment 為該筆資料是否屬於實驗組，sb、bs 與 bb 則是其於沒門檻/有門檻遊戲當中的大小隊分配情形')
        tr_sb_p.add_run('\n- 實驗組的 wtp 為原始數據，至於控制組則是以三種方法計算，包含兩個遊戲之總和、取最大值以及取最小值')
        
        space = self.document.add_paragraph()
        space.add_run(' ')        
        self.document.add_heading('控制組的 wtp 為兩個遊戲的總和', level = 2)
        self.document.add_heading('線性迴歸模型結果', level = 3)
        
        self.add_table(sum_reg)
        self.document.add_heading('F 檢定', level = 3)
        self.add_table(sum_f)
        
        space = self.document.add_paragraph()
        space.add_run(' ')
        self.document.add_heading('控制組的 wtp 為個遊戲當中的最大值', level = 2)
        self.document.add_heading('線性迴歸模型結果', level = 3)
        self.add_table(max_reg)
        self.document.add_heading('F 檢定', level = 3)
        self.add_table(max_f)

        space = self.document.add_paragraph()
        space.add_run(' ')

        self.document.add_heading('控制組的 wtp 為兩個遊戲當中的最小值', level = 2)
        self.document.add_heading('線性迴歸模型結果', level = 3)
        self.add_table(min_reg)
        self.document.add_heading('F 檢定', level = 3)
        self.add_table(min_f)
        
    def save_document(self):
        today = date.today()
        day = today.strftime("%B %d, %Y")
        fileName = day + "_report.docx"
        self.document.save(fileName)
        return fileName
